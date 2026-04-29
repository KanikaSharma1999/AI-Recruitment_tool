import streamlit as st
import pandas as pd
import time
from utils import read_file, rank_resumes, clean_text
from cohere_feedback import get_resume_feedback
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import io
import os
import re
import spacy

# Load spaCy model for NLP name extraction
nlp = spacy.load("en_core_web_sm")

def extract_candidate_details(raw_text, filename):
    """
    Extract name, email, and phone from resume text.

    Name extraction uses a STRICT 5-strategy multi-retry system with:
      - Banned keyword rejection (section headings, job titles)
      - Format validation (2–3 words, Title Case, alpha-only, ≤40 chars)
      - Email cross-check (at least one name token must appear in email username)
      - Safe email-based fallback when all strategies fail
    """

    # ── STEP 1: EMAIL EXTRACTION (SOURCE OF TRUTH) ────────────────────────────
    email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
    email = email_match.group(0) if email_match else "Not Found"

    # Derive a clean username token for cross-checking
    email_username = ""
    if email != "Not Found":
        raw_username = email.split("@")[0]
        email_username = re.sub(r'[^a-z]', '', raw_username.lower())  # strip digits/dots

    # ── STEP 2: PHONE EXTRACTION ──────────────────────────────────────────────
    phone_match = re.search(r'(\+?\d{1,3}[.\-\s]?)?\d{10}', raw_text)
    phone = phone_match.group(0) if phone_match else "Not Found"

    # ── HELPERS ───────────────────────────────────────────────────────────────

    # Words that must NEVER appear in a candidate name
    BANNED_WORDS = {
        "experience", "summary", "profile", "objective", "overview",
        "manager", "engineer", "developer", "sales", "analyst",
        "consultant", "director", "executive", "officer", "lead",
        "resume", "cv", "curriculum", "vitae", "contact", "details",
        "education", "skills", "projects", "achievements", "references",
        "professional", "personal", "technical", "work", "employment",
        "history", "background", "information", "about", "page",
    }

    def is_valid_name(candidate: str) -> bool:
        """Return True only when candidate passes ALL strict rules."""
        candidate = candidate.strip()
        if not candidate:
            return False
        words = candidate.split()
        # Rule 1: must be 2–3 words
        if not (2 <= len(words) <= 3):
            return False
        # Rule 2: total length ≤ 40
        if len(candidate) > 40:
            return False
        # Rule 3: no digits
        if re.search(r'\d', candidate):
            return False
        # Rule 4: no @ symbol
        if '@' in candidate:
            return False
        # Rule 5: every word must be alpha-only
        if not all(w.isalpha() for w in words):
            return False
        # Rule 6: each word must start with uppercase (Title Case)
        if not all(w[0].isupper() for w in words):
            return False
        # Rule 7: must NOT be entirely uppercase (e.g. "WORK EXPERIENCE")
        if candidate == candidate.upper():
            return False
        # Rule 8: must NOT be entirely lowercase
        if candidate == candidate.lower():
            return False
        # Rule 9: no banned keyword
        if any(w.lower() in BANNED_WORDS for w in words):
            return False
        return True

    def email_cross_check(candidate: str) -> bool:
        """At least one word from the candidate name must appear in email username."""
        if not email_username:
            return True  # no email → skip check (don't penalise)
        name_tokens = [w.lower() for w in candidate.split()]
        return any(token in email_username or email_username.startswith(token[:3])
                   for token in name_tokens if len(token) >= 3)

    def try_name(candidate: str) -> str | None:
        """Validate + cross-check; return cleaned name or None."""
        candidate = ' '.join(candidate.split())           # normalise whitespace
        candidate = ' '.join(w.capitalize() for w in candidate.split())
        if is_valid_name(candidate) and email_cross_check(candidate):
            return candidate
        return None

    # ── STEP 3–5: MULTI-STRATEGY EXTRACTION WITH RETRY ───────────────────────
    lines = raw_text.split('\n')
    first_30 = [l.strip() for l in lines[:30] if l.strip()]

    name = None

    # --- TRY 1: Very first non-empty line ---
    if not name and first_30:
        name = try_name(first_30[0])

    # --- TRY 2: First 3 non-empty lines ---
    if not name:
        for line in first_30[:3]:
            result = try_name(line)
            if result:
                name = result
                break

    # --- TRY 3: Lines near/before the email address ---
    if not name and email != "Not Found":
        email_line_idx = None
        for idx, raw_line in enumerate(lines[:30]):
            if email in raw_line:
                email_line_idx = idx
                break
        if email_line_idx is not None:
            # Check up to 5 lines before the email line
            for raw_line in lines[max(0, email_line_idx - 5): email_line_idx]:
                result = try_name(raw_line.strip())
                if result:
                    name = result
                    break

    # --- TRY 4: spaCy PERSON entity extraction ---
    if not name:
        doc = nlp(raw_text[:2000])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                result = try_name(ent.text)
                if result:
                    name = result
                    break

    # --- TRY 5: Full scan of first 30 lines for any valid name pattern ---
    if not name:
        for line in first_30:
            # also try sub-phrases (e.g. "Name: John Smith")
            # strip common label prefixes
            cleaned = re.sub(r'^(name\s*[:|-]?\s*)', '', line, flags=re.IGNORECASE).strip()
            result = try_name(cleaned)
            if result:
                name = result
                break

    # ── STEP 6: FINAL FALLBACK – derive from email username ──────────────────
    if not name:
        if email != "Not Found":
            raw_username = email.split("@")[0]
            # Remove digits and common separators
            clean_username = re.sub(r'[\d._\-]+', ' ', raw_username).strip()
            parts = [p.capitalize() for p in clean_username.split() if p]
            if parts:
                name = ' '.join(parts[:2])   # use at most 2 tokens from username
            else:
                # username was entirely digits – use just the stem capitalised
                stem = re.sub(r'\d', '', raw_username).capitalize()
                name = stem if stem else os.path.splitext(filename)[0].replace('_', ' ').strip()
        else:
            # Last resort: cleaned filename
            name = os.path.splitext(filename)[0].replace('_', ' ').strip()

    return {
        "name": name,
        "email": email,
        "phone": phone,
    }


def generate_report(job_text, ranking, resumes, candidate_details):
    report = "AI Resume Ranking Report\n\n"

    report += "JOB DESCRIPTION\n"
    report += job_text + "\n\n"

    report += "RANKING SECTION\n"
    for i, (filename, score, candidate_name) in enumerate(ranking, 1):
        report += f"{i}. {candidate_name} — Score: {score:.2f}\n"

    report += "\nALL CANDIDATES FEEDBACK\n\n"

    for i, (filename, score, candidate_name) in enumerate(ranking, 1):
        details = candidate_details[filename]
        report += f"Candidate Name: {candidate_name}\n"
        report += f"Score: {score:.2f}\n"
        report += f"Email: {details['email']}\n"
        report += f"Phone: {details['phone']}\n\n"
        feedback = get_resume_feedback(resumes[filename], job_text)
        report += feedback + "\n\n"

    return report

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = parse_xml(r'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{}"><w:r><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr><w:t>{}</w:t></w:r></w:hyperlink>'.format(r_id, text))
    paragraph._p.append(hyperlink)
    return hyperlink

def generate_word_report(job_text, ranking, resumes, saved_paths, candidate_details):
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Resume Ranking Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
    
    doc.add_paragraph()  # Spacing
    
    # Job Description Section
    job_heading = doc.add_heading('Job Description', level=1)
    job_run = job_heading.runs[0]
    job_run.font.size = Pt(16)
    job_run.font.bold = True
    job_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark Blue
    
    job_para = doc.add_paragraph(job_text)
    job_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Candidate Ranking Section
    ranking_heading = doc.add_heading('Candidate Ranking', level=1)
    ranking_run = ranking_heading.runs[0]
    ranking_run.font.size = Pt(16)
    ranking_run.font.bold = True
    ranking_run.font.color.rgb = RGBColor(102, 0, 102)  # Purple
    
    for i, (filename, score, candidate_name) in enumerate(ranking, 1):
        para = doc.add_paragraph()
        para.add_run(f"{i}. ").bold = True
        para.add_run(f"{candidate_name}").bold = True
        para.add_run(" — Score: ")
        score_run = para.add_run(f"{score:.2f}")
        score_run.font.highlight_color = 7  # Yellow highlight
    
    doc.add_paragraph()  # Spacing
    
    # All Candidate Analysis Section
    analysis_heading = doc.add_heading('All Candidate Analysis', level=1)
    analysis_run = analysis_heading.runs[0]
    analysis_run.font.size = Pt(16)
    analysis_run.font.bold = True
    analysis_run.font.color.rgb = RGBColor(0, 102, 0)  # Green
    
    for i, (filename, score, candidate_name) in enumerate(ranking, 1):
        details = candidate_details[filename]
        
        # Candidate header
        candidate_para = doc.add_paragraph()
        candidate_run = candidate_para.add_run(f"Candidate {i}: {candidate_name}")
        candidate_run.font.bold = True
        candidate_run.font.size = Pt(14)
        
        # Score and contact details
        doc.add_paragraph(f"Score: {score:.2f}")
        doc.add_paragraph(f"Email: {details['email']}")
        doc.add_paragraph(f"Phone: {details['phone']}")
        doc.add_paragraph(f"Resume File: {filename}")
        
        # Clickable link to resume
        link_para = doc.add_paragraph()
        file_path = saved_paths[filename]
        file_url = "file:///" + file_path.replace("\\", "/")
        add_hyperlink(link_para, "🔗 View Resume", file_url)
        
        doc.add_paragraph()  # Spacing
        
        # Get feedback
        feedback = get_resume_feedback(resumes[filename], job_text)
        
        # Parse and format feedback as plain bullet points
        feedback_lines = feedback.split('\n')
        for line in feedback_lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('• '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line:
                doc.add_paragraph(line, style='List Bullet')
        
        doc.add_paragraph()  # Spacing between candidates
    
    # Convert to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

# ------------------- PAGE CONFIG -------------------
st.set_page_config(page_title="AI Resume Ranker", layout="wide")

# ------------------- CUSTOM CSS -------------------
st.markdown("""
<style>
.main {
    background-color: #f8f9fc;
}
.card {
    padding: 20px;
    border-radius: 12px;
    background-color: white;
    box-shadow: 0px 4px 12px rgba(0,0,0,0.05);
    margin-bottom: 20px;
}
</style>
""", unsafe_allow_html=True)

# ------------------- TITLE -------------------
st.title("📝 AI Resume Ranker & Feedback System")
st.markdown("Upload job description and resumes to rank candidates using AI.")

# ------------------- FILE UPLOAD -------------------
st.markdown("### 📂 Upload Files")

job_file = st.file_uploader("Upload Job Description (.txt)", type=['txt'])
resume_files = st.file_uploader(
    "Upload Resumes (PDF or TXT)",
    type=['pdf', 'txt'],
    accept_multiple_files=True
)

# ------------------- MAIN LOGIC -------------------
if job_file and resume_files:
    # STEP 1: SEPARATE RAW TEXT AND CLEAN TEXT
    job_raw_text = read_file(job_file)
    job_text = clean_text(job_raw_text)  # For ranking

    # Store both raw and cleaned text
    raw_resumes = {}
    resumes = {}  # cleaned for ranking

    for file in resume_files:
        raw_text = read_file(file)
        cleaned_text = clean_text(raw_text)
        raw_resumes[file.name] = raw_text
        resumes[file.name] = cleaned_text

    # STEP 2: EXTRACT CANDIDATE DETAILS FROM RAW TEXT
    candidate_details = {}
    for filename, raw_text in raw_resumes.items():
        candidate_details[filename] = extract_candidate_details(raw_text, filename)

    # STEP 3: SAVE UPLOADED RESUMES LOCALLY FOR CLICKABLE LINKS
    resumes_dir = "uploaded_resumes"
    os.makedirs(resumes_dir, exist_ok=True)

    saved_paths = {}
    for file in resume_files:
        file_path = os.path.join(resumes_dir, file.name)
        with open(file_path, "wb") as f:
            f.write(file.getbuffer())
        saved_paths[file.name] = os.path.abspath(file_path)

    if st.button("🚀 Run Analysis"):

        # STEP 4: RANK RESUMES USING CLEANED TEXT
        ranking = rank_resumes(job_text, resumes)

        # STEP 5: ENHANCE RANKING WITH CANDIDATE NAMES
        enhanced_ranking = []
        for filename, score in ranking:
            candidate_name = candidate_details[filename]["name"]
            enhanced_ranking.append((filename, score, candidate_name))

        # ------------------- TABS -------------------
        tab1, tab2, tab3 = st.tabs(["📊 Ranking", "📈 Visualization", "📢 Feedback"])

        # ------------------- TAB 1: RANKING -------------------
        with tab1:
            st.subheader("📊 Resume Ranking")

            results_data = []

            for rank, (filename, score, candidate_name) in enumerate(enhanced_ranking, 1):
                details = candidate_details[filename]

                file_path = saved_paths[filename]
                file_url = "file:///" + file_path.replace("\\", "/")

                st.markdown(f"""
                <div class="card">
                    <h4>🏆 Rank {rank}: {candidate_name}</h4>
                    <p><b>Similarity Score:</b> {score:.2f}</p>
                    <p><b>Email:</b> <a href="mailto:{details['email']}">{details['email']}</a></p>
                    <p><b>Phone:</b> {details['phone']}</p>
                    <p><b>Resume:</b> <a href="{file_url}" target="_blank">🔗 View Resume</a></p>
                </div>
                """, unsafe_allow_html=True)

                st.markdown("---")

                results_data.append({
                    "Resume": candidate_name,
                    "Score": score
                })

            # 📥 Download CSV
            df = pd.DataFrame(results_data)
            csv = df.to_csv(index=False).encode('utf-8')

            st.download_button(
                label="📥 Download Results",
                data=csv,
                file_name="resume_results.csv",
                mime="text/csv"
            )

            # Generate and download full report
            report_text = generate_report(job_text, enhanced_ranking, raw_resumes, candidate_details)
            st.download_button(
                label="📥 Download Full Report",
                data=report_text,
                file_name="resume_report.txt",
                mime="text/plain"
            )

            # Generate and download professional Word report
            word_report_bytes = generate_word_report(job_text, enhanced_ranking, raw_resumes, saved_paths, candidate_details)
            st.download_button(
                label="📄 Download Professional Report",
                data=word_report_bytes,
                file_name="AI_Resume_Report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # ------------------- TAB 2: VISUALIZATION -------------------
        with tab2:
            st.subheader("📈 Score Visualization")

            chart_data = {candidate_name: score for _, score, candidate_name in enhanced_ranking}
            st.bar_chart(chart_data)

        # ------------------- TAB 3: FEEDBACK -------------------
        with tab3:
            st.subheader("📢 AI Feedback (All Candidates)")
            st.info("Short AI evaluation for every uploaded resume based on the job description.")

            progress_bar = st.progress(0, text="Generating feedback for all resumes...")
            total = len(enhanced_ranking)

            for i, (filename, score, candidate_name) in enumerate(enhanced_ranking, 1):
                details = candidate_details[filename]

                # Update progress
                progress_bar.progress(i / total, text=f"Processing {i} of {total}: {candidate_name}")

                with st.expander(f"🏆 Rank {i}: {candidate_name} — Score: {score:.2f}", expanded=(i == 1)):

                    # Candidate info row
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown(f"**👤 Name:** {candidate_name}")
                        st.markdown(f"**📧 Email:** [{details['email']}](mailto:{details['email']})")
                    with col2:
                        st.markdown(f"**📞 Phone:** {details['phone']}")
                        st.markdown(f"**📊 Score:** `{score:.2f}`")

                    st.divider()

                    # Generate and display feedback
                    with st.spinner("Generating feedback..."):
                        feedback = get_resume_feedback(raw_resumes[filename], job_raw_text)
                        time.sleep(0.5)  # Avoid API rate limiting

                    st.markdown(feedback)

            progress_bar.progress(1.0, text="✅ Feedback generated for all candidates!")

else:
    st.info("Upload both Job Description and Resumes to proceed.")